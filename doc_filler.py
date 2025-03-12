import os
import re
import threading
import pandas as pd
from tkinter import *
from tkinterdnd2 import DND_FILES, TkinterDnD
from docx import Document
from PyPDF2 import PdfReader
import requests
from datetime import datetime

class DocFillerApp:
    def __init__(self, master):
        self.master = master
        master.title("智能文档填空系统 v1.0")
        
        # API密钥输入
        self.api_frame = LabelFrame(master, text="DeepSeek API 配置")
        self.api_frame.pack(fill=X, padx=10, pady=5)
        
        self.api_label = Label(self.api_frame, text="API Key:")
        self.api_label.pack(side=LEFT, padx=5)
        self.api_entry = Entry(self.api_frame, width=40, show="*")
        self.api_entry.pack(side=LEFT, fill=X, expand=True)
        
        # 文档上传区
        self.upload_frame = Frame(master)
        self.upload_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        
        # 分析文档上传
        self.doc_frame = LabelFrame(self.upload_frame, text="拖拽上传待分析文档 (支持多文件)")
        self.doc_frame.pack(side=LEFT, fill=BOTH, expand=True, padx=5)
        self.doc_list = Listbox(self.doc_frame, width=40, height=10)
        self.doc_list.pack(fill=BOTH, expand=True, padx=5, pady=5)
        self.doc_list.drop_target_register(DND_FILES)
        self.doc_list.dnd_bind('<<Drop>>', self.add_docs)
        
        # 模板上传
        self.tpl_frame = LabelFrame(self.upload_frame, text="拖拽上传填空模板")
        self.tpl_frame.pack(side=LEFT, fill=BOTH, expand=True, padx=5)
        self.tpl_list = Listbox(self.tpl_frame, width=40, height=10)
        self.tpl_list.pack(fill=BOTH, expand=True, padx=5, pady=5)
        self.tpl_list.drop_target_register(DND_FILES)
        self.tpl_list.dnd_bind('<<Drop>>', self.add_tpl)
        
        # 控制按钮
        self.btn_frame = Frame(master)
        self.btn_frame.pack(fill=X, padx=10, pady=5)
        
        self.clear_btn = Button(self.btn_frame, text="清空列表", command=self.clear_lists)
        self.clear_btn.pack(side=LEFT, padx=5)
        
        self.start_btn = Button(self.btn_frame, text="开始处理", command=self.start_processing)
        self.start_btn.pack(side=RIGHT, padx=5)
        
        # 日志显示
        self.log_frame = LabelFrame(master, text="处理日志")
        self.log_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        
        self.log_text = Text(self.log_frame, height=10)
        self.scroll = Scrollbar(self.log_frame, command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=self.scroll.set)
        self.scroll.pack(side=RIGHT, fill=Y)
        self.log_text.pack(fill=BOTH, expand=True)
        
        # 状态栏
        self.status = Label(master, text="准备就绪", bd=1, relief=SUNKEN, anchor=W)
        self.status.pack(side=BOTTOM, fill=X)
        
        # 初始化变量
        self.docs = []
        self.templates = []
        self.lock = threading.Lock()
        self.running = False

    def add_docs(self, event):
        files = self.parse_dropped_files(event.data)
        new_files = [f for f in files if f not in self.docs]
        self.docs.extend(new_files)
        self.update_listbox(self.doc_list, new_files)
        self.log(f"添加 {len(new_files)} 个分析文档")
        
    def add_tpl(self, event):
        files = self.parse_dropped_files(event.data)
        if len(files) > 1:
            self.log("错误：只能上传一个模板文件")
            return
        self.templates = files
        self.update_listbox(self.tpl_list, files)
        self.log("模板文件已更新")
    
    def parse_dropped_files(self, data):
        files = []
        for item in data.split():
            filepath = item.replace("{", "").replace("}", "")
            if os.path.exists(filepath):
                files.append(filepath)
        return files
    
    def start_processing(self):
        if not self.validate_inputs():
            return
        self.running = True
        threading.Thread(target=self.process_files).start()
    
    def validate_inputs(self):
        if not self.docs:
            self.log("错误：请先上传分析文档")
            return False
        if not self.templates:
            self.log("错误：请先上传模板文件")
            return False
        if not self.api_entry.get().strip():
            self.log("错误：请输入API密钥")
            return False
        return True
    
    def process_files(self):
        try:
            template_path = self.templates[0]
            combined_content = []
            
            # 读取所有分析文档内容
            for doc_path in self.docs:
                content = self.read_file_content(doc_path)
                if content:
                    combined_content.append(f"文档内容：{os.path.basename(doc_path)}\n{content}")
            
            # 合并内容并调用API
            final_content = "\n\n".join(combined_content)
            self.log("正在调用DeepSeek API...")
            answers = self.call_deepseek_api(final_content)
            
            if not answers:
                self.log("API调用失败，请检查日志")
                return
                
            # 填充模板并保存
            self.fill_template(template_path, answers)
            self.log("处理完成！")
            
        except Exception as e:
            self.log(f"处理出错：{str(e)}")
        finally:
            self.running = False

    def read_file_content(self, file_path):
        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.pdf':
                reader = PdfReader(file_path)
                return "\n".join([page.extract_text() for page in reader.pages])
            elif ext in ('.docx', '.doc'):
                doc = Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            elif ext in ('.xlsx', '.xls'):
                # 使用openpyxl引擎增强Excel兼容性
                df = pd.read_excel(file_path, engine='openpyxl')
                return df.to_string()
            else:
                self.log(f"不支持的文件格式：{ext}")
                return None
        except Exception as e:
            self.log(f"读取文件失败：{os.path.basename(file_path)} - {str(e)}")
            return None

    def call_deepseek_api(self, content):
        api_key = os.getenv("DEEPSEEK_API_KEY") or self.api_entry.get().strip()
        if not api_key:
            self.log("错误：API密钥未配置，请设置环境变量DEEPSEEK_API_KEY或在界面输入")
            return None
            
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "deepseek-chat",
            "messages": [{
                "role": "user",
                "content": f"请根据以下内容填写模板中的空白：\n{content}\n请直接返回填写好的内容，不要添加额外说明。"
            }],
            "temperature": 0.3
        }
        try:
            response = requests.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=30
            )
            response.raise_for_status()
            return response.json()["choices"][0]["message"]["content"]
        except Exception as e:
            self.log(f"API调用失败：{str(e)}")
            return None

    def fill_template(self, template_path, answers):
        try:
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"模板文件不存在: {template_path}")
                
            doc = Document(template_path)
            found_placeholder = False
            
            # 处理段落
            for p in doc.paragraphs:
                if '___' in p.text:
                    p.text = re.sub(r'_{3,}', answers, p.text)
                    found_placeholder = True
                    
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if '___' in cell.text:
                            cell.text = re.sub(r'_{3,}', answers, cell.text)
                            found_placeholder = True
            
            if not found_placeholder:
                self.log("警告：未在模板中找到任何填空占位符（___）")
                return

            output_dir = os.path.join(os.path.expanduser("~"), "Documents")
            os.makedirs(output_dir, exist_ok=True)
            output_name = f"{os.path.splitext(os.path.basename(template_path))[0]}_已回答_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
            output_path = os.path.join(output_dir, output_name)
            
            doc.save(output_path)
            self.log(f"成功生成结果文件：{output_path}")
            return output_path
        except PermissionError as e:
            self.log(f"文件权限错误：请关闭正在使用的模板文件 - {str(e)}")
        except FileNotFoundError as e:
            self.log(f"文件未找到：{str(e)}")
        except Exception as e:
            self.log(f"生成结果文件失败：{str(e)}")
            raise

    def log(self, message):
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert(END, f"[{timestamp}] {message}\n")
        self.log_text.see(END)
        self.status.config(text=message)
    
    def update_listbox(self, listbox, items):
        listbox.delete(0, END)
        for item in items:
            listbox.insert(END, os.path.basename(item))
    
    def clear_lists(self):
        self.docs = []
        self.templates = []
        self.update_listbox(self.doc_list, [])
        self.update_listbox(self.tpl_list, [])
        self.log("已清空所有文件列表")

if __name__ == "__main__":
    root = TkinterDnD.Tk()
    app = DocFillerApp(root)
    root.geometry("800x600")
    root.mainloop()
